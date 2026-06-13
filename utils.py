# -*- coding: utf-8 -*-
"""
文档工具函数
提供模板数据提取、纯文本提取和格式应用功能
"""

import re
import logging
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


# ==============================
#  辅助函数（序列化）
# ==============================

def _serialize_font_size(size_emu):
    if size_emu is None:
        return None
    try:
        return round(size_emu / 12700, 1)
    except (TypeError, AttributeError):
        return None


def _serialize_color(color_obj):
    if color_obj is None:
        return None
    try:
        if hasattr(color_obj, 'rgb'):
            return str(color_obj.rgb)
        return str(color_obj)
    except (TypeError, AttributeError):
        return None


def _serialize_alignment(alignment):
    if alignment is None:
        return None
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: '左对齐',
        WD_ALIGN_PARAGRAPH.CENTER: '居中',
        WD_ALIGN_PARAGRAPH.RIGHT: '右对齐',
        WD_ALIGN_PARAGRAPH.JUSTIFY: '两端对齐',
    }
    return mapping.get(alignment, str(alignment))


# ==============================
#  格式名称 与 python-docx 常量 的映射
# ==============================

def _parse_alignment(text):
    """将中文对齐描述转换为 WD_ALIGN_PARAGRAPH 枚举"""
    mapping = {
        '左对齐': WD_ALIGN_PARAGRAPH.LEFT,
        '居中': WD_ALIGN_PARAGRAPH.CENTER,
        '右对齐': WD_ALIGN_PARAGRAPH.RIGHT,
        '两端对齐': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(text)


def _parse_font_size_cn(text):
    """将中文字号转换为磅值，如 '二号' -> 22"""
    mapping = {
        '初号': 42, '小初': 36, '一号': 26, '小一': 24,
        '二号': 22, '小二': 18, '三号': 16, '小三': 15,
        '四号': 14, '小四': 12, '五号': 10.5, '小五': 9,
        '六号': 7.5, '小六': 6.5, '七号': 5.5, '八号': 5,
    }
    return mapping.get(text)


def _get_header_footer_data(section):
    result = {}
    if section.header and section.header.paragraphs:
        header_paras = [p.text.strip() for p in section.header.paragraphs if p.text.strip()]
        if header_paras:
            result['header'] = header_paras
    if section.footer and section.footer.paragraphs:
        footer_paras = [p.text.strip() for p in section.footer.paragraphs if p.text.strip()]
        if footer_paras:
            result['footer'] = footer_paras
    return result


# ==============================
#  标签 → 模板规则键名 映射
# ==============================

def _map_label_to_rule_key(label):
    """将段落语义标签映射到 template_rules['style_rules'] 中的键名"""
    mapping = {
        '论文大标题': '论文大标题',
        '作者信息': '作者信息',
        '摘要标题': '摘要标题',
        '摘要正文': '摘要正文',
        '关键词': '关键词',
        '一级标题': '一级标题',
        '二级标题': '二级标题',
        '三级标题': '三级标题',
        '正文段落': '正文',
        '图表标题': '图表标题',
        '参考文献标题': '参考文献标题',
        '参考文献条目': '参考文献条目',
        '致谢': '正文',
        '附录标题': '一级标题',
        '附录内容': '正文',
        '页眉': '页眉',
        '页脚': '页脚',
    }
    return mapping.get(label, '正文')


# ===== GB/T 7713.1-2006 国标后备规则（当 template_rules 中没有时使用）=====
GB_FALLBACK_RULES = {
    '论文大标题': {
        'chinese_font': {'name': '黑体', 'size': '二号', 'bold': True},
        'western_font': {'name': 'Times New Roman', 'size': '二号', 'bold': True},
        'paragraph': {'alignment': '居中', 'space_before': '1行', 'space_after': '1行'},
    },
    '作者信息': {
        'chinese_font': {'name': '楷体', 'size': '四号', 'bold': False},
        'western_font': {'name': 'Times New Roman', 'size': '四号', 'bold': False},
        'paragraph': {'alignment': '居中'},
    },
    '摘要标题': {
        'chinese_font': {'name': '黑体', 'size': '小四', 'bold': True},
        'western_font': {'name': 'Times New Roman', 'size': '小四', 'bold': True},
        'paragraph': {'alignment': '两端对齐'},
    },
    '摘要正文': {
        'chinese_font': {'name': '宋体', 'size': '小四', 'bold': False},
        'western_font': {'name': 'Times New Roman', 'size': '小四', 'bold': False},
        'paragraph': {'alignment': '两端对齐', 'first_line_indent': '首行缩进2字符'},
    },
    '关键词': {
        'chinese_font': {'name': '黑体', 'size': '小四', 'bold': True},
        'western_font': {'name': 'Times New Roman', 'size': '小四', 'bold': True},
        'paragraph': {'alignment': '两端对齐'},
    },
    '一级标题': {
        'chinese_font': {'name': '黑体', 'size': '小三', 'bold': True},
        'western_font': {'name': 'Times New Roman', 'size': '小三', 'bold': True},
        'paragraph': {'alignment': '居中', 'space_before': '0.5行', 'space_after': '0.5行'},
    },
    '二级标题': {
        'chinese_font': {'name': '黑体', 'size': '四号', 'bold': True},
        'western_font': {'name': 'Times New Roman', 'size': '四号', 'bold': True},
        'paragraph': {'alignment': '左对齐', 'space_before': '0.5行', 'space_after': '0.5行'},
    },
    '三级标题': {
        'chinese_font': {'name': '黑体', 'size': '小四', 'bold': True},
        'western_font': {'name': 'Times New Roman', 'size': '小四', 'bold': True},
        'paragraph': {'alignment': '左对齐', 'space_before': '0.5行', 'space_after': '0.5行'},
    },
    '正文': {
        'chinese_font': {'name': '宋体', 'size': '小四', 'bold': False},
        'western_font': {'name': 'Times New Roman', 'size': '小四', 'bold': False},
        'paragraph': {'alignment': '两端对齐', 'first_line_indent': '首行缩进2字符', 'line_spacing': '1.5倍'},
    },
    '图表标题': {
        'chinese_font': {'name': '黑体', 'size': '五号', 'bold': False},
        'western_font': {'name': 'Times New Roman', 'size': '五号', 'bold': False},
        'paragraph': {'alignment': '居中', 'space_before': '0.5行', 'space_after': '0.5行'},
    },
    '参考文献标题': {
        'chinese_font': {'name': '黑体', 'size': '小三', 'bold': True},
        'western_font': {'name': 'Times New Roman', 'size': '小三', 'bold': True},
        'paragraph': {'alignment': '居中'},
    },
    '参考文献条目': {
        'chinese_font': {'name': '宋体', 'size': '五号', 'bold': False},
        'western_font': {'name': 'Times New Roman', 'size': '五号', 'bold': False},
        'paragraph': {'alignment': '两端对齐'},
    },
    '页眉': {
        'chinese_font': {'name': '宋体', 'size': '五号', 'bold': False},
        'western_font': {'name': 'Times New Roman', 'size': '五号', 'bold': False},
        'paragraph': {'alignment': '居中'},
    },
    '页脚': {
        'chinese_font': {'name': 'Times New Roman', 'size': '小五', 'bold': False},
        'western_font': {'name': 'Times New Roman', 'size': '小五', 'bold': False},
        'paragraph': {'alignment': '居中'},
    },
}


def _resolve_rule(rule_key, style_rules):
    """
    从 style_rules 中查找规则，找不到则使用国标后备规则。
    """
    # 先从模板规则中找
    rule = style_rules.get(rule_key) if style_rules else None
    if rule:
        return rule
    # 从国标后备中找
    gb_rule = GB_FALLBACK_RULES.get(rule_key)
    if gb_rule:
        logger.info("使用国标后备规则: %s", rule_key)
        return dict(gb_rule)
    # 最终后备：国标正文规则
    logger.warning("未找到规则 %s，使用国标正文后备", rule_key)
    return dict(GB_FALLBACK_RULES.get('正文', {}))


# ==============================
#  核心格式应用函数（Bug修复版）
# ==============================

def _apply_run_font_clean(run, chinese_font, western_font):
    """
    【Bug修复】彻底清除 run 原有字体格式后重新设置。

    修复: "不该换行的地方换行" 问题。
    方法: 不操作原有的 XML，而是直接设置 python-docx 的字体属性，
    确保不会引入多余的空 run 或换行符。

    通过 XML 设置 w:eastAsia 属性确保中文字体正确显示。
    """
    if run is None:
        return

    ch_name = (chinese_font or {}).get('name', '宋体')
    w_name = (western_font or {}).get('name', 'Times New Roman')
    ch_size = (chinese_font or {}).get('size', '小四')
    ch_bold = (chinese_font or {}).get('bold', False)

    # 设置英文字体
    run.font.name = w_name
    # 通过 XML 设置中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), ch_name)
    rFonts.set(qn('w:ascii'), w_name)
    rFonts.set(qn('w:hAnsi'), w_name)

    # 设置字号（磅值）
    size_pt = _parse_font_size_cn(ch_size)
    if size_pt:
        run.font.size = Pt(size_pt)

    # 设置粗体（明确设置 True 或 False，不保留 None）
    run.font.bold = bool(ch_bold)


def _apply_paragraph_format_clean(paragraph, rule, label):
    """
    【强硬修复】应用段落级格式。

    ===== 对齐方式强制规则 =====
    - 标签含"标题"/"大标题" (不含"参考文献标题") → CENTER
    - 标签含"摘要"或"关键词" → JUSTIFY
    - 标签含"正文" → JUSTIFY
    - 标签含"参考文献条目" → JUSTIFY
    - 其他 → 从规则读取，若规则没有则用 JUSTIFY

    ===== 首行缩进强制规则 =====
    - 正文/摘要内容 → Cm(0.74) ≈ 2字符
    - 标题类 → None (清空)
    """
    pf = paragraph.paragraph_format
    align = None  # 最终对齐方式

    # 【强制对齐逻辑】
    if '标题' in label and '参考文献标题' not in label:
        align = WD_ALIGN_PARAGRAPH.CENTER
    elif '大标题' in label:
        align = WD_ALIGN_PARAGRAPH.CENTER
    elif '摘要' in label or '关键词' in label:
        align = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif '正文' in label:
        align = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif '参考文献条目' in label:
        align = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        # 从规则读取
        para_fmt = rule.get('paragraph', {})
        align_text = para_fmt.get('alignment')
        if align_text:
            align = _parse_alignment(align_text)
        if align is None:
            align = WD_ALIGN_PARAGRAPH.JUSTIFY

    if align is not None:
        paragraph.alignment = align
    print(f"[Format] 段落 ({label}): 对齐方式设置为 {paragraph.alignment} (CENTER=1, JUSTIFY=3)")

    # 行距
    ls_text = (rule.get('paragraph', {}) or {}).get('line_spacing')
    if ls_text:
        try:
            s = str(ls_text)
            if '倍' in s:
                pf.line_spacing = float(s.replace('倍', ''))
            else:
                val = float(s)
                pf.line_spacing = val if val < 10 else Pt(val)
        except (ValueError, TypeError):
            pass

    # 段前间距
    for key in ['space_before', 'space_after']:
        val_text = (rule.get('paragraph', {}) or {}).get(key)
        if val_text:
            try:
                s = str(val_text)
                if '行' in s:
                    setattr(pf, key, Pt(float(s.replace('行', '')) * 12))
                elif '磅' in s:
                    setattr(pf, key, Pt(float(s.replace('磅', ''))))
                else:
                    setattr(pf, key, Pt(float(s)))
            except (ValueError, TypeError):
                pass

    # 【强制首行缩进】
    is_heading = ('标题' in label and '参考文献标题' not in label) or '大标题' in label
    is_body = ('正文' in label or '摘要' in label)

    if is_heading:
        pf.first_line_indent = Cm(0)
        print(f"[Format] 段落 ({label}): 首行缩进已清除")
    elif is_body:
        pf.first_line_indent = Cm(0.74)
        print(f"[Format] 段落 ({label}): 首行缩进设为 0.74cm")
    else:
        # 从规则读取
        indent = (rule.get('paragraph', {}) or {}).get('first_line_indent')
        if indent:
            s = str(indent)
            try:
                if '字符' in s:
                    n = re.search(r'(\d+)', s)
                    if n:
                        pf.first_line_indent = Cm(float(n.group(1)) * 0.37)
                elif 'cm' in s.lower():
                    cm_v = float(re.search(r'([\d.]+)', s).group(1))
                    pf.first_line_indent = Cm(cm_v)
                else:
                    pf.first_line_indent = Cm(float(s) * 0.035)
            except (ValueError, TypeError, AttributeError):
                pass


def apply_format(template_path, paper_path, output_path, labels=None, template_rules=None):
    """
    根据 labels 和 template_rules 将格式规则应用到论文文档。

    【段创建铁律】每个段落先清空所有 run，再添加唯一一个新 run。
    【国标后备】缺失规则自动用 GB/T 7713.1-2006 填补。
    【对齐强制】标题→居中，正文→两端对齐。

    参数:
        template_path: 模板路径
        paper_path: 原论文路径
        output_path: 输出路径
        labels: {段落索引: 标签名} 字典
        template_rules: analyze_template 返回的格式规范表
    """
    logger.info("开始应用格式 - 论文: %s, 输出: %s", paper_path, output_path)

    doc = Document(paper_path)
    style_rules = (template_rules or {}).get('style_rules', {})

    if not labels:
        labels = {}

    applied_count = 0
    skipped_empty = 0

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        label = labels.get(str(idx), labels.get(idx, '正文段落'))
        rule_key = _map_label_to_rule_key(label)

        # 【国标后备】_resolve_rule 会先查模板规则，找不到则用国标
        rule = _resolve_rule(rule_key, style_rules)
        if not rule:
            continue

        chinese_font = rule.get('chinese_font', {})
        western_font = rule.get('western_font', {})

        # ===== 【段创建铁律】清空所有 run，重建唯一 run =====
        # 1. 保存原文本
        original_text = paragraph.text

        # 2. 清空所有 run
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

        # 3. 添加唯一一个干净的新 run
        new_run = paragraph.add_run(original_text)

        # 4. 在这个唯一 run 上设置字体
        _apply_run_font_clean(new_run, chinese_font, western_font)

        # 【强硬修复】应用段落格式（含对齐修正）
        _apply_paragraph_format_clean(paragraph, rule, label)

        if original_text.strip():
            applied_count += 1
        else:
            skipped_empty += 1

    doc.save(output_path)
    logger.info(
        "格式应用完成 - 处理 %d 个非空段落, 跳过 %d 个空段落, 输出: %s",
        applied_count, skipped_empty, output_path
    )


# ==============================
#  提取函数（与之前一致）
# ==============================

def extract_template_data(docx_path):
    """从模板文档中提取每个段落的完整格式信息"""
    doc = Document(docx_path)
    result = {'paragraphs': [], 'headers': [], 'footers': [], 'page_margins': {}}

    section = doc.sections[0] if doc.sections else None
    if section:
        margins = section.page_margins if hasattr(section, 'page_margins') else section
        result['page_margins'] = {
            'top': _serialize_font_size(getattr(margins, 'top', None)),
            'bottom': _serialize_font_size(getattr(margins, 'bottom', None)),
            'left': _serialize_font_size(getattr(margins, 'left', None)),
            'right': _serialize_font_size(getattr(margins, 'right', None)),
        }

    for section in doc.sections:
        hf = _get_header_footer_data(section)
        if 'header' in hf:
            result['headers'].extend(hf['header'])
        if 'footer' in hf:
            result['footers'].extend(hf['footer'])

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else None
        run = para.runs[0] if para.runs else None
        paragraph_info = {
            'text': text[:50] if text else '',
            'text_length': len(text),
            'style_name': style_name,
            'font': {
                'name': run.font.name if run else None,
                'size': _serialize_font_size(run.font.size if run else None),
                'bold': run.font.bold if run else None,
                'italic': run.font.italic if run else None,
                'color': _serialize_color(run.font.color.rgb if run and run.font.color and run.font.color.rgb else None),
            },
            'paragraph_format': {
                'alignment': _serialize_alignment(para.alignment),
                'line_spacing': _serialize_font_size(para.paragraph_format.line_spacing),
                'space_before': _serialize_font_size(para.paragraph_format.space_before),
                'space_after': _serialize_font_size(para.paragraph_format.space_after),
                'first_line_indent': _serialize_font_size(para.paragraph_format.first_line_indent),
            },
        }
        result['paragraphs'].append(paragraph_info)

    logger.info("模板数据提取完成 - 共 %d 个段落", len(result['paragraphs']))
    return result


def extract_docx_text(docx_path):
    """提取文档的纯文本内容，保留段落结构"""
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return '\n'.join(paragraphs)
